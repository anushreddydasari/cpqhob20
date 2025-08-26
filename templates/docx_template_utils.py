from docx import Document
from typing import Dict


def replace_placeholders_in_docx_bytes(docx_bytes: bytes, variables: Dict[str, str]) -> bytes:
    """Replace {{placeholders}} in a DOCX (bytes) and return modified bytes.

    Robust against run-splitting by reconstructing paragraph text and rewriting runs.
    Tables are handled by iterating cells and applying the same logic.
    """
    from io import BytesIO

    input_buffer = BytesIO(docx_bytes)
    document = Document(input_buffer)

    def _replace_all(text: str) -> str:
        new_text = text
        for key, value in variables.items():
            placeholder = f"{{{{{key}}}}}"
            new_text = new_text.replace(placeholder, str(value))
        return new_text

    def replace_in_paragraph(paragraph):
        # Join all runs' text, replace placeholders, then set as a single run
        full_text = ''.join(run.text or '' for run in paragraph.runs)
        replaced = _replace_all(full_text)
        if replaced == full_text:
            # No change: keep as-is
            return
        # Clear all runs
        for _ in range(len(paragraph.runs)):
            paragraph.runs[0].clear()
            paragraph._element.remove(paragraph.runs[0]._element)
        # Add one run with replaced text
        paragraph.add_run(replaced)

    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph)

    for table in document.tables:
        replace_in_table(table)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


