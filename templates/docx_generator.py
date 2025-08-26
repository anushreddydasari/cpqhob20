from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re
from datetime import datetime
import base64
import requests
from PIL import Image

class DocxGenerator:
    def __init__(self):
        self.document = None
    
    def create_document(self, template_content, template_data):
        """Create a DOCX document from template content and data"""
        try:
            # Create new document
            self.document = Document()
            
            # Set up document styles
            self._setup_document_styles()
            
            # Sanitize/normalize HTML, then replace placeholders
            cleaned_html = self._preprocess_html(template_content)
            processed_content = self._replace_placeholders(cleaned_html, template_data)
            
            # Convert HTML content to DOCX
            self._convert_html_to_docx(processed_content)
            
            return True
        except Exception as e:
            print(f"Error creating DOCX document: {str(e)}")
            return False

    def _preprocess_html(self, html_content: str) -> str:
        """Simplify editor HTML so our basic converter handles it nicely.
        - Strip <span ...> wrappers and inline styles
        - Convert <mark> to plain text
        - Normalize <br> to newlines
        - Collapse multiple spaces and &nbsp;
        """
        try:
            content = html_content or ''
            # Remove span wrappers but keep inner text
            content = re.sub(r'</?span[^>]*>', '', content, flags=re.IGNORECASE)
            # Replace mark with its inner text
            content = re.sub(r'</?mark[^>]*>', '', content, flags=re.IGNORECASE)
            # Normalize <br> to newline
            content = re.sub(r'<br\s*/?>', '\n', content, flags=re.IGNORECASE)
            # Remove style attrs inside headings/paragraphs we handle
            content = re.sub(r'<(h[1-6]|p)[^>]*>', lambda m: f"<{m.group(1)}>", content, flags=re.IGNORECASE)
            # Strip stray leading characters (pipes, bullets, zero-width) after tag open
            content = re.sub(r'(<h[1-6]>)[\s\|\u2022\u200B\u200C\u200D\ufeff]+', r'\1', content)
            content = re.sub(r'(<p>)[\s\|\u2022\u200B\u200C\u200D\ufeff]+', r'\1', content)
            # Convert &nbsp; -> space
            content = content.replace('&nbsp;', ' ')
            # Trim trailing spaces per line
            content = '\n'.join([ln.rstrip() for ln in content.split('\n')])
            return content
        except Exception:
            return html_content
    
    def _setup_document_styles(self):
        """Set up document styles and formatting"""
        # Set default font
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Create heading styles
        heading1_style = self.document.styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.font.name = 'Calibri'
        heading1_style.font.size = Pt(16)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = None  # Default color
        
        heading2_style = self.document.styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.font.name = 'Calibri'
        heading2_style.font.size = Pt(14)
        heading2_style.font.bold = True
        
        heading3_style = self.document.styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        heading3_style.font.name = 'Calibri'
        heading3_style.font.size = Pt(12)
        heading3_style.font.bold = True
    
    def _replace_placeholders(self, content, data):
        """Replace template placeholders with actual data"""
        try:
            
            
            # Replace all other placeholders in the format {{placeholder_name}}
            placeholder_pattern = r'\{\{(\w+)\}\}'
            
            def replace_placeholder(match):
                placeholder_name = match.group(1)
                return str(data.get(placeholder_name, f'[{placeholder_name}]'))
            
            processed_content = re.sub(placeholder_pattern, replace_placeholder, content)
            return processed_content
        except Exception as e:
            print(f"Error replacing placeholders: {str(e)}")
            return content


    
    def _convert_html_to_docx(self, html_content):
        """Convert HTML content to DOCX format"""
        try:
            # Tokenize by blocks to support multi-line tags
            content = html_content
            block_pattern = re.compile(
                r'(<h1>.*?</h1>|<h2>.*?</h2>|<h3>.*?</h3>|<p>.*?</p>|<ul>.*?</ul>|<ol>.*?</ol>|<blockquote>.*?</blockquote>|<table[^>]*?>.*?</table>|<img[^>]*>)',
                re.IGNORECASE | re.DOTALL
            )
            pos = 0
            for m in block_pattern.finditer(content):
                # Text before the block
                pre = content[pos:m.start()].strip()
                if pre:
                    self._emit_plain_text(pre)
                block = m.group(1)
                self._emit_block(block)
                pos = m.end()
            # Tail text
            tail = content[pos:].strip()
            if tail:
                self._emit_plain_text(tail)
        
        except Exception as e:
            print(f"Error converting HTML to DOCX: {str(e)}")
            # Fallback: add as plain text
            self.document.add_paragraph(html_content)

    def _emit_plain_text(self, text: str):
        # Remove residual tags and add as paragraph
        clean = re.sub(r'<[^>]+>', '', text)
        clean = clean.strip()
        if clean:
            self.document.add_paragraph(clean)

    def _emit_block(self, block_html: str):
        block = block_html.strip()
        lower = block.lower()
        if lower.startswith('<h1>'):
            # Emit any inline images within the heading
            self._emit_inline_images(block)
            text = self._strip_tags(re.sub(r'<img[^>]*>', '', block, flags=re.IGNORECASE))
            p = self.document.add_paragraph(text, style='Custom Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        if lower.startswith('<h2>'):
            self._emit_inline_images(block)
            text = self._strip_tags(re.sub(r'<img[^>]*>', '', block, flags=re.IGNORECASE))
            p = self.document.add_paragraph(text, style='Custom Heading 2')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        if lower.startswith('<h3>'):
            self._emit_inline_images(block)
            text = self._strip_tags(re.sub(r'<img[^>]*>', '', block, flags=re.IGNORECASE))
            p = self.document.add_paragraph(text, style='Custom Heading 3')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        if lower.startswith('<p>'):
            # Emit images inside the paragraph, then add remaining text
            self._emit_inline_images(block)
            no_imgs = re.sub(r'<img[^>]*>', '', block, flags=re.IGNORECASE)
            text = self._strip_tags(no_imgs)
            if text:
                self.document.add_paragraph(text)
            return
        if lower.startswith('<img'):
            self._process_image(block)
            return
        if lower.startswith('<ul>') or lower.startswith('<ol>'):
            items = re.findall(r'<li>(.*?)</li>', block, flags=re.IGNORECASE | re.DOTALL)
            for it in items:
                self.document.add_paragraph(self._strip_tags(it))
            return
        if lower.startswith('<blockquote>'):
            self._emit_inline_images(block)
            text = self._strip_tags(re.sub(r'<img[^>]*>', '', block, flags=re.IGNORECASE))
            if text:
                para = self.document.add_paragraph(text)
                para.style = 'Quote'
            return
        if lower.startswith('<table'):
            self._process_table(block)
            return
        # Fallback to plain
        self._emit_plain_text(block)
    
    def _extract_text(self, html_line, tag):
        """Extract text content from HTML tag"""
        try:
            # Simple text extraction
            start_tag = f'<{tag}>'
            end_tag = f'</{tag}>'
            
            if start_tag in html_line and end_tag in html_line:
                start_idx = html_line.find(start_tag) + len(start_tag)
                end_idx = html_line.find(end_tag)
                return html_line[start_idx:end_idx].strip()
            return html_line
        except Exception as e:
            print(f"Error extracting text from {tag}: {str(e)}")
            return html_line
    
    def _process_list(self, list_html):
        """Process HTML list elements"""
        try:
            # Extract list items
            items = re.findall(r'<li>(.*?)</li>', list_html, re.DOTALL)
            
            for item in items:
                text = item.strip()
                if text:
                    paragraph = self.document.add_paragraph()
                    paragraph.add_run('• ').bold = True
                    paragraph.add_run(text)
        
        except Exception as e:
            print(f"Error processing list: {str(e)}")
    
    def _process_table(self, table_html):
        """Process HTML table elements"""
        try:
            # Extract table rows
            rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL | re.IGNORECASE)

            if rows:
                # Determine max columns across rows
                # Extract cells; keep empty strings if missing
                all_cells = [re.findall(r'<t(?:d|h)[^>]*>(.*?)</t(?:d|h)>', r, re.DOTALL | re.IGNORECASE) for r in rows]
                max_cols = max((len(c) for c in all_cells), default=1)

                table = self.document.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'
                try:
                    table.autofit = False
                except Exception:
                    pass

                # Compute equal column widths to approximate HTML width
                try:
                    section = self.document.sections[0]
                    available_width_in = float((section.page_width - section.left_margin - section.right_margin)) / 914400.0
                    col_width_in = max(0.5, available_width_in / max_cols)
                    from docx.shared import Inches as _Inches
                    for j in range(max_cols):
                        for cell in table.columns[j].cells:
                            cell.width = _Inches(col_width_in)
                except Exception:
                    pass

                for i, cells in enumerate(all_cells):
                    for j in range(max_cols):
                        cell = table.rows[i].cells[j]
                        # Remove existing paragraph content safely
                        for p in list(cell.paragraphs):
                            p._element.getparent().remove(p._element)

                        text_html = cells[j] if j < len(cells) else ''

                        # Emit images first (as separate paragraphs in the cell)
                        for img in re.finditer(r'<img[^>]*>', text_html, flags=re.IGNORECASE):
                            # Insert picture into cell by adding a paragraph run
                            para_for_img = cell.add_paragraph()
                            self._process_image(img.group(0))

                        # Convert inner HTML line breaks to paragraph breaks
                        normalized = re.sub(r'<br\s*/?>', '\n', text_html, flags=re.IGNORECASE)
                        # Remove image tags for text portion
                        normalized = re.sub(r'<img[^>]*>', '', normalized, flags=re.IGNORECASE)

                        # Split on paragraphs if any (<p>...</p>) inside cell
                        inner_ps = re.findall(r'<p[^>]*>(.*?)</p>', normalized, flags=re.IGNORECASE | re.DOTALL)
                        if inner_ps:
                            for seg in inner_ps:
                                clean = self._strip_tags(seg)
                                clean = re.sub(r'\s+', ' ', clean).strip()
                                if clean:
                                    cell.add_paragraph(clean)
                        else:
                            clean = self._strip_tags(normalized)
                            clean = re.sub(r'\s+', ' ', clean).strip()
                            if clean:
                                cell.add_paragraph(clean)

                # Style header row if first row has any <th>
                try:
                    header_is_th = bool(re.search(r'<th', rows[0], flags=re.IGNORECASE))
                    if header_is_th:
                        header_row = table.rows[0]
                        for hdr_cell in header_row.cells:
                            # Bold + center paragraphs
                            if not hdr_cell.paragraphs:
                                hdr_cell.add_paragraph('')
                            for p in hdr_cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for r in p.runs:
                                    r.bold = True
                            # Shading (light gray)
                            try:
                                tcPr = hdr_cell._tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:fill'), 'F1F3F5')
                                shd.set(qn('w:val'), 'clear')
                                tcPr.append(shd)
                            except Exception:
                                pass
                except Exception:
                    pass
        
        except Exception as e:
            print(f"Error processing table: {str(e)}")

    def _emit_inline_images(self, html: str):
        """Emit any <img> tags contained within a larger HTML block (e.g., inside <p>, headings, table cells)."""
        try:
            for m in re.finditer(r'<img[^>]*>', html, flags=re.IGNORECASE):
                self._process_image(m.group(0))
        except Exception:
            pass

    def _parse_img_attrs(self, img_html: str):
        """Extract src, width, height from an <img> tag string."""
        try:
            src_match = re.search(r'src\s*=\s*"([^"]+)"', img_html, flags=re.IGNORECASE)
            if not src_match:
                src_match = re.search(r"src\s*=\s*'([^']+)'", img_html, flags=re.IGNORECASE)
            src = src_match.group(1) if src_match else None

            width = None
            height = None
            # Try width/height attributes
            w_match = re.search(r'width\s*=\s*"(\d+)"', img_html, flags=re.IGNORECASE)
            if not w_match:
                w_match = re.search(r"width\s*=\s*'(\d+)'", img_html, flags=re.IGNORECASE)
            h_match = re.search(r'height\s*=\s*"(\d+)"', img_html, flags=re.IGNORECASE)
            if not h_match:
                h_match = re.search(r"height\s*=\s*'(\d+)'", img_html, flags=re.IGNORECASE)
            if w_match:
                width = int(w_match.group(1))
            if h_match:
                height = int(h_match.group(1))

            # Try inline style width/height (e.g., style="width:200px;height:80px")
            style_match = re.search(r'style\s*=\s*"([^"]+)"', img_html, flags=re.IGNORECASE)
            if style_match:
                style = style_match.group(1)
                sw = re.search(r'width\s*:\s*(\d+)px', style, flags=re.IGNORECASE)
                sh = re.search(r'height\s*:\s*(\d+)px', style, flags=re.IGNORECASE)
                if sw:
                    width = int(sw.group(1))
                if sh:
                    height = int(sh.group(1))

            return src, width, height
        except Exception:
            return None, None, None

    def _process_image(self, img_html: str):
        """Process <img> tag: supports http(s) URLs and data URIs. Optional width/height attrs in px.
        """
        try:
            src, width_px, height_px = self._parse_img_attrs(img_html)
            if not src:
                return

            image_bytes = None
            if src.startswith('data:image/'):
                # data URI
                b64_match = re.search(r'base64,([A-Za-z0-9+/=]+)', src)
                if b64_match:
                    image_bytes = base64.b64decode(b64_match.group(1))
            elif src.startswith('http://') or src.startswith('https://'):
                # Some CDNs block requests without UA; include a common UA header
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
                }
                resp = requests.get(src, timeout=10, headers=headers)
                if resp.status_code == 200:
                    image_bytes = resp.content
            else:
                # treat as local file path or app-served uploads (/uploads/..)
                try:
                    local_path = src
                    if src.startswith('/uploads/'):
                        import os
                        # Map to UPLOAD_FOLDER
                        base = os.path.join(os.getcwd(), 'uploaded_docs')
                        local_path = os.path.join(base, src.split('/uploads/', 1)[1])
                    with open(local_path, 'rb') as f:
                        image_bytes = f.read()
                except Exception:
                    image_bytes = None

            if not image_bytes:
                return

            from io import BytesIO
            stream = BytesIO(image_bytes)

            # Normalize to a docx-friendly format if needed (e.g., WebP → PNG)
            try:
                img = Image.open(BytesIO(image_bytes))
                fmt = (img.format or '').upper()
                if fmt not in ('PNG', 'JPEG', 'JPG', 'GIF', 'BMP', 'TIFF'):
                    buf = BytesIO()
                    img.convert('RGBA').save(buf, format='PNG')
                    buf.seek(0)
                    stream = buf
            except Exception:
                # If PIL can't open, fall back to raw bytes (may fail silently)
                pass

            # Determine target size and auto-fit to page width
            # Read intrinsic image size (pixels)
            img_w_px = None
            img_h_px = None
            try:
                meta_img = Image.open(BytesIO(image_bytes))
                img_w_px, img_h_px = meta_img.size
            except Exception:
                pass

            # Compute requested size in inches (default from intrinsic size)
            width_in = width_px / 96.0 if width_px else (img_w_px / 96.0 if img_w_px else None)
            height_in = height_px / 96.0 if height_px else (img_h_px / 96.0 if img_h_px else None)

            # If only height provided, derive width by aspect ratio
            if width_in is None and height_in and img_w_px and img_h_px:
                width_in = (img_w_px / img_h_px) * height_in
            # If neither provided, and no metadata, fallback to 2 inches
            if width_in is None:
                width_in = 2.0

            # Clamp to available page width (preserve aspect ratio)
            try:
                section = self.document.sections[0]
                # python-docx uses EMUs; convert to inches (1 inch = 914400 EMUs)
                available_width_in = float((section.page_width - section.left_margin - section.right_margin)) / 914400.0
                if width_in > available_width_in:
                    scale = available_width_in / width_in
                    width_in = available_width_in
                    if height_in:
                        height_in = height_in * scale
            except Exception:
                pass

            kwargs = {'width': Inches(width_in)}
            # Do not set height to preserve aspect ratio unless explicitly needed
            # If height was explicitly provided and width not adjusted above we could set it, but width is preferred

            self.document.add_picture(stream, **kwargs)
        except Exception as e:
            print(f"Error processing image: {str(e)}")

    def _strip_tags(self, html_fragment: str) -> str:
        try:
            # Remove all tags and reduce whitespace
            text = re.sub(r'<[^>]+>', '', html_fragment)
            text = text.replace('&nbsp;', ' ')
            return ' '.join(text.split())
        except Exception:
            return html_fragment
    
    def add_header(self, title, subtitle=None):
        """Add document header"""
        try:
            # Add title
            title_paragraph = self.document.add_paragraph(title)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_paragraph.style = 'Custom Heading 1'
            
            # Add subtitle if provided
            if subtitle:
                subtitle_paragraph = self.document.add_paragraph(subtitle)
                subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_paragraph.style = 'Custom Heading 2'
            
            # Add separator
            self.document.add_paragraph()
        
        except Exception as e:
            print(f"Error adding header: {str(e)}")
    
    def add_footer(self, company_name, generated_date=None):
        """Add document footer"""
        try:
            if not generated_date:
                generated_date = datetime.now().strftime('%B %d, %Y')
            
            # Add separator
            self.document.add_paragraph()
            
            # Add footer information
            footer_paragraph = self.document.add_paragraph()
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_paragraph.add_run(f'Generated by {company_name} on {generated_date}')
            footer_paragraph.style.font.size = Pt(9)
            footer_paragraph.style.font.italic = True
        
        except Exception as e:
            print(f"Error adding footer: {str(e)}")
    
    def save_document(self, file_path):
        """Save the document to file"""
        try:
            if self.document:
                self.document.save(file_path)
                return True
            return False
        except Exception as e:
            print(f"Error saving document: {str(e)}")
            return False
    
    def get_document_bytes(self):
        """Get document as bytes for download"""
        try:
            if self.document:
                from io import BytesIO
                buffer = BytesIO()
                self.document.save(buffer)
                buffer.seek(0)
                return buffer.getvalue()
            return None
        except Exception as e:
            print(f"Error getting document bytes: {str(e)}")
            return None

def generate_agreement_docx(template_content, template_data, output_path=None):
    """Generate DOCX agreement from template"""
    try:
        generator = DocxGenerator()
        
        # Create document
        success = generator.create_document(template_content, template_data)
        if not success:
            return False, "Failed to create document"
        
        # Add header
        generator.add_header(
            f"Agreement - {template_data.get('service_type', 'Services')}",
            f"Client: {template_data.get('client_name', 'N/A')}"
        )
        
        # Add footer
        generator.add_footer(
            template_data.get('company_name', 'Your Company'),
            template_data.get('generation_date')
        )
        
        # Save or return bytes
        if output_path:
            success = generator.save_document(output_path)
            return success, "Document saved successfully" if success else "Failed to save document"
        else:
            doc_bytes = generator.get_document_bytes()
            return True, doc_bytes if doc_bytes else "Failed to get document bytes"
    
    except Exception as e:
        return False, f"Error generating DOCX: {str(e)}"
